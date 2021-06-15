import json
import requests
import io
import os
import urllib
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.shared import Pt

get_headers = {
    'Authorization': 'Bearer '
    }

questions_url = 'https://api.airtable.com/v0/appl09q3uKHeKNUnq/Export_to_word?view=Grid%20view'
questions = requests.get(questions_url, headers=get_headers)
questions_data = questions.json()

no_questions = len(questions_data['records'])

document = Document('template.docx')
font = document.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(12)

intro = document.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro_font = intro.add_run('PART I')
intro_font.bold = True

rule = document.add_paragraph()
rule.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
rule.add_run(f'For each question from 1 to {no_questions}, four options are given. One of them is the correct answer. Make your choice (1, 2, 3 or 4). Shade the correct oval (1, 2, 3 or 4) on the Optical Answer Sheet.                                                                (24 x 2 marks)')

divider = document.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
divider.add_run('_' * 64)

for i in questions_data['records']:
    question = i['fields']['Question']
    
    if '<diagram>' in question:
        pictures = i['fields']['Attachments']
        picture_url = []
        for picture in pictures:
            picture_url.append(picture['url'])
        question_list = question.split('<diagram>')
        diagram_count = question.count('<diagram>')
        count = 0
        p = document.add_paragraph(style='List Number')
        for question_part in question_list:
            p_text = p.add_run(question_part)
            if diagram_count != count:
                io_url = io.BytesIO(urllib.request.urlopen(picture_url[count]).read())
                p_text.add_picture(io_url)
            count += 1
        bracket = document.add_paragraph()
        bracket.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = bracket.add_run(f'(\t)')

        if question == questions_data['records'][no_questions - 1]['fields']['Question']:
            end = document.add_paragraph()
            end.alignment = WD_ALIGN_PARAGRAPH.CENTER
            end_font = end.add_run('END OF BOOKLET A')
            end_font.bold = True
        else:
#             run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
            
    else:
        p = document.add_paragraph(question, style='List Number')
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

document.save('test.docx')
os.system('start test.docx')